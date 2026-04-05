"""
RxMonitor Fetcher Microservice
FastAPI service responsible for:
1. Fetching artifacts from all 5 real source types
2. Extracting text with pdfplumber / BeautifulSoup
3. Computing SHA-256 hashes for change detection
4. Uploading raw artifacts to Supabase Storage
5. Returning structured text + metadata to the Next.js pipeline
"""

import hashlib
import io
import os
import tempfile
from datetime import datetime
from typing import Any
from urllib.parse import urlparse

import httpx
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from dotenv import load_dotenv
from pathlib import Path
from fastapi import FastAPI, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from supabase import create_client, Client
from tenacity import retry, stop_after_attempt, wait_exponential

# Load .env.local from project root (Next.js convention) with fallback to .env
_root = Path(__file__).parent.parent
load_dotenv(_root / ".env.local")
load_dotenv(_root / ".env", override=False)

app = FastAPI(title="RxMonitor Fetcher", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[os.getenv("NEXT_PUBLIC_SITE_URL", "http://localhost:3000")],
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

SUPABASE_URL = os.getenv("NEXT_PUBLIC_SUPABASE_URL", "")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")
FETCHER_API_SECRET = os.getenv("FETCHER_API_SECRET", "")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

# ---------------------------------------------------------------------------
# Response Models
# ---------------------------------------------------------------------------

class PageSection(BaseModel):
    section_title: str | None
    page_number: int | None
    text: str

class FetchResult(BaseModel):
    source_id: str
    content_hash: str
    storage_path: str
    sections: list[PageSection]
    total_pages: int | None
    fetch_method: str
    fetched_at: str
    changed: bool  # True if hash differs from last known version

# ---------------------------------------------------------------------------
# Auth guard
# ---------------------------------------------------------------------------

def verify_secret(x_api_secret: str = Header(...)):
    if x_api_secret != FETCHER_API_SECRET:
        raise HTTPException(status_code=401, detail="Invalid API secret")

# ---------------------------------------------------------------------------
# Shared HTTP client with retries
# ---------------------------------------------------------------------------

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (compatible; RxMonitor/1.0; "
        "+https://rxmonitor.app/bot)"
    ),
    "Accept": "text/html,application/xhtml+xml,application/pdf,*/*",
}

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
async def fetch_url(url: str, timeout: int = 30) -> httpx.Response:
    async with httpx.AsyncClient(follow_redirects=True, timeout=timeout) as client:
        resp = await client.get(url, headers=HEADERS)
        resp.raise_for_status()
        return resp

# ---------------------------------------------------------------------------
# PDF text extraction with pdfplumber (coordinate-aware, table-friendly)
# ---------------------------------------------------------------------------

def extract_pdf_sections(pdf_bytes: bytes) -> tuple[list[PageSection], int]:
    sections: list[PageSection] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        total = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            # Also extract tables as text
            for table in page.extract_tables():
                rows = ["\t".join(str(c) for c in row if c) for row in table if row]
                text += "\n" + "\n".join(rows)
            if text.strip():
                sections.append(PageSection(
                    section_title=None,
                    page_number=page.page_number,
                    text=text.strip(),
                ))
    return sections, total


def extract_pdf_tables(pdf_bytes: bytes) -> list[list[str]]:
    """Extract all table rows from a PDF (for Priority Health MDL pure-table format)."""
    all_rows: list[list[str]] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    if row and any(cell for cell in row if cell):
                        all_rows.append([str(c or "") for c in row])
    return all_rows


def extract_docx_sections(docx_bytes: bytes) -> list[PageSection]:
    """Extract sections from a .docx file using python-docx."""
    sections: list[PageSection] = []
    doc = DocxDocument(io.BytesIO(docx_bytes))

    current_title: str | None = None
    current_text: list[str] = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if "heading" in style_name or style_name.startswith("title"):
            # Flush previous section
            if current_text:
                sections.append(PageSection(
                    section_title=current_title,
                    page_number=None,
                    text=" ".join(current_text),
                ))
                current_text = []
            current_title = text
        else:
            current_text.append(text)

    # Flush last section
    if current_text:
        sections.append(PageSection(
            section_title=current_title,
            page_number=None,
            text=" ".join(current_text),
        ))

    # Also extract text from tables in the docx
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            sections.append(PageSection(
                section_title="Table",
                page_number=None,
                text="\n".join(rows_text),
            ))

    return sections or [PageSection(section_title=None, page_number=None, text="(empty document)")]

# ---------------------------------------------------------------------------
# HTML extraction with BeautifulSoup
# ---------------------------------------------------------------------------

def extract_html_sections(html: str, base_url: str) -> list[PageSection]:
    soup = BeautifulSoup(html, "lxml")
    # Remove nav, footer, scripts, styles
    for tag in soup(["nav", "footer", "script", "style", "header", "aside"]):
        tag.decompose()

    sections: list[PageSection] = []
    current_title: str | None = None
    current_text: list[str] = []

    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th"]):
        if element.name in ("h1", "h2", "h3", "h4"):
            # Flush previous section
            if current_text:
                sections.append(PageSection(
                    section_title=current_title,
                    page_number=None,
                    text=" ".join(current_text).strip(),
                ))
                current_text = []
            current_title = element.get_text(strip=True)
        else:
            t = element.get_text(strip=True)
            if t:
                current_text.append(t)

    if current_text:
        sections.append(PageSection(
            section_title=current_title,
            page_number=None,
            text=" ".join(current_text).strip(),
        ))

    return sections or [PageSection(section_title=None, page_number=None,
                                    text=soup.get_text(separator=" ", strip=True))]

# ---------------------------------------------------------------------------
# CSV / JSON extraction (for CMS QHP formulary data)
# ---------------------------------------------------------------------------

def extract_csv_sections(raw: str, source_name: str) -> list[PageSection]:
    """Convert CSV rows to text blocks for embedding/extraction."""
    lines = raw.splitlines()
    if not lines:
        return []
    header = lines[0]
    chunks: list[str] = []
    # Group 50 rows per section to create coherent chunks
    batch: list[str] = [header]
    for i, line in enumerate(lines[1:], 1):
        batch.append(line)
        if len(batch) >= 51:
            chunks.append("\n".join(batch))
            batch = [header]
    if len(batch) > 1:
        chunks.append("\n".join(batch))

    return [
        PageSection(section_title=f"{source_name} rows {i*50+1}-{(i+1)*50}",
                    page_number=None, text=chunk)
        for i, chunk in enumerate(chunks)
    ]

def extract_json_sections(data: Any, source_name: str) -> list[PageSection]:
    """Convert JSON API response to text blocks."""
    if isinstance(data, dict) and "results" in data:
        items = data["results"]
    elif isinstance(data, list):
        items = data
    else:
        return [PageSection(section_title=source_name, page_number=None, text=str(data))]

    # Group 20 items per section
    sections = []
    for i in range(0, len(items), 20):
        batch = items[i:i+20]
        text = "\n".join(str(item) for item in batch)
        sections.append(PageSection(
            section_title=f"{source_name} items {i+1}-{i+len(batch)}",
            page_number=None,
            text=text,
        ))
    return sections

# ---------------------------------------------------------------------------
# Supabase storage
# ---------------------------------------------------------------------------

def upload_to_storage(content: bytes, source_id: str, content_hash: str, ext: str) -> str:
    path = f"artifacts/{source_id}/{content_hash[:16]}.{ext}"
    supabase.storage.from_("artifacts").upload(
        path, content,
        file_options={"content-type": "application/octet-stream", "upsert": "true"},
    )
    return path

def get_last_hash(source_id: str) -> str | None:
    result = (
        supabase.table("artifact_versions")
        .select("content_hash")
        .eq("source_id", source_id)
        .order("fetched_at", desc=True)
        .limit(1)
        .execute()
    )
    rows = result.data
    return rows[0]["content_hash"] if rows else None

# ---------------------------------------------------------------------------
# Per-source fetch logic
# ---------------------------------------------------------------------------

async def fetch_pdf_source(source: dict) -> tuple[bytes, list[PageSection], int | None]:
    resp = await fetch_url(source["fetch_url"])
    pdf_bytes = resp.content
    sections, total = extract_pdf_sections(pdf_bytes)
    return pdf_bytes, sections, total

async def fetch_html_source(source: dict) -> tuple[bytes, list[PageSection], None]:
    resp = await fetch_url(source["fetch_url"])
    html = resp.text
    sections = extract_html_sections(html, source["fetch_url"])
    return html.encode(), sections, None

async def fetch_api_source(source: dict) -> tuple[bytes, list[PageSection], None]:
    resp = await fetch_url(source["fetch_url"])
    ct = resp.headers.get("content-type", "")
    raw = resp.content

    if "json" in ct:
        data = resp.json()
        sections = extract_json_sections(data, source["name"])
    elif "csv" in ct or source["fetch_url"].endswith(".csv"):
        sections = extract_csv_sections(resp.text, source["name"])
    else:
        sections = [PageSection(section_title=source["name"], page_number=None,
                                text=resp.text[:50000])]
    return raw, sections, None

# ---------------------------------------------------------------------------
# CMS Medicare Coverage DB — special handling for structured API
# ---------------------------------------------------------------------------

async def fetch_medicare_cms(source: dict) -> tuple[bytes, list[PageSection], None]:
    """
    Fetches NCD/LCD articles from CMS Medicare Coverage Database API.
    The API returns a list of coverage documents; for each, we fetch the
    full article text.
    """
    # Step 1: search for GLP-1 articles
    search_url = source["fetch_url"]
    resp = await fetch_url(search_url)
    data = resp.json()

    articles = data.get("items") or data.get("results") or (data if isinstance(data, list) else [])
    all_text = f"CMS Medicare Coverage Database — GLP-1 Policies\nFetched: {datetime.utcnow().isoformat()}\n\n"
    sections: list[PageSection] = []

    for article in articles[:10]:  # Limit to 10 for rate limiting
        title = article.get("title") or article.get("Title", "Untitled")
        doc_type = article.get("DocumentType") or article.get("articleType", "")
        eff_date = article.get("effectiveDate") or article.get("EffectiveDate", "")

        # Try to get full text from the article URL if available
        article_url = article.get("articleUrl") or article.get("URL") or ""
        article_text = f"Title: {title}\nType: {doc_type}\nEffective Date: {eff_date}\n"

        if article_url and article_url.startswith("http"):
            try:
                art_resp = await fetch_url(article_url, timeout=15)
                if "html" in art_resp.headers.get("content-type", ""):
                    html_sections = extract_html_sections(art_resp.text, article_url)
                    article_text += "\n".join(s.text for s in html_sections[:5])
            except Exception:
                article_text += "(full text unavailable)"

        sections.append(PageSection(
            section_title=title,
            page_number=None,
            text=article_text,
        ))
        all_text += f"\n---\n{article_text}"

    return all_text.encode(), sections, None

# ---------------------------------------------------------------------------
# VA National Formulary — download full formulary document
# ---------------------------------------------------------------------------

async def fetch_va_formulary(source: dict) -> tuple[bytes, list[PageSection], int | None]:
    # VA publishes the National Formulary as a downloadable PDF link on the page
    resp = await fetch_url(source["fetch_url"])
    soup = BeautifulSoup(resp.text, "lxml")

    # Find the PDF link on the page
    pdf_link = None
    for a in soup.find_all("a", href=True):
        href: str = a["href"]
        if "formulary" in href.lower() and href.endswith(".pdf"):
            pdf_link = href if href.startswith("http") else "https://www.pbm.va.gov" + href
            break

    if pdf_link:
        pdf_resp = await fetch_url(pdf_link)
        pdf_bytes = pdf_resp.content
        sections, total = extract_pdf_sections(pdf_bytes)
        return pdf_bytes, sections, total
    else:
        # Fall back to HTML extraction
        html_sections = extract_html_sections(resp.text, source["fetch_url"])
        return resp.content, html_sections, None

# ---------------------------------------------------------------------------
# Main /fetch endpoint
# ---------------------------------------------------------------------------

FETCH_DISPATCH = {
    "pdf": fetch_pdf_source,
    "html": fetch_html_source,
    "api": fetch_api_source,
    "csv": fetch_api_source,
}

SPECIAL_HANDLERS = {
    "medicare": fetch_medicare_cms,
    "va_tricare": fetch_va_formulary,
}

@app.get("/fetch/{source_id}", response_model=FetchResult)
async def fetch_source(
    source_id: str,
    x_api_secret: str = Header(...),
):
    verify_secret(x_api_secret)

    # Load source config from Supabase
    result = supabase.table("sources").select("*").eq("id", source_id).single().execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Source not found")
    source = result.data

    if not source["active"]:
        raise HTTPException(status_code=400, detail="Source is not active")

    # Choose fetcher
    plan_type = source["plan_type"]
    fetch_method = source["fetch_method"]

    if plan_type in SPECIAL_HANDLERS:
        handler = SPECIAL_HANDLERS[plan_type]
    else:
        handler = FETCH_DISPATCH.get(fetch_method, fetch_html_source)

    try:
        raw_bytes, sections, total_pages = await handler(source)
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"Upstream error: {e.response.status_code}")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Fetch failed: {str(e)}")

    # Change detection
    content_hash = hashlib.sha256(raw_bytes).hexdigest()
    last_hash = get_last_hash(source_id)
    changed = last_hash != content_hash

    # Determine extension
    ext_map = {"pdf": "pdf", "html": "html", "csv": "csv", "api": "json"}
    ext = ext_map.get(fetch_method, "bin")

    # Upload to Supabase Storage only if changed (or no prior version)
    if changed:
        storage_path = upload_to_storage(raw_bytes, source_id, content_hash, ext)
    else:
        # Return existing path
        existing = (
            supabase.table("artifact_versions")
            .select("storage_path")
            .eq("source_id", source_id)
            .order("fetched_at", desc=True)
            .limit(1)
            .execute()
        )
        storage_path = existing.data[0]["storage_path"] if existing.data else ""

    return FetchResult(
        source_id=source_id,
        content_hash=content_hash,
        storage_path=storage_path,
        sections=sections,
        total_pages=total_pages,
        fetch_method=fetch_method,
        fetched_at=datetime.utcnow().isoformat(),
        changed=changed,
    )

class PipelineRequest(BaseModel):
    source_id: str
    artifact_version_id: str
    user_id: str | None = None

@app.post("/pipeline/run")
async def run_pipeline(
    req: PipelineRequest,
    x_api_secret: str = Header(...),
):
    """
    Trigger the full LangGraph pipeline for a source.
    Expects the artifact_version to already be inserted into DB by the caller
    (fetch-check/route.ts does this before calling here).

    Flow: load artifact from Storage → chunk → embed → extract → evaluate → persist draft
    NOTE: Does NOT re-fetch the source. The caller already did the fetch + change detection
    + DB insert. Re-fetching would always return changed=False (hash already in DB) and
    cause the pipeline to never run.
    """
    verify_secret(x_api_secret)

    from pipeline.graph import pipeline_graph

    # Load source metadata
    source_result = supabase.table("sources").select("*").eq("id", req.source_id).single().execute()
    if not source_result.data:
        raise HTTPException(status_code=404, detail="Source not found")
    source = source_result.data

    # Load the artifact version that was already inserted by the caller
    version_result = (
        supabase.table("artifact_versions")
        .select("*")
        .eq("id", req.artifact_version_id)
        .single()
        .execute()
    )
    if not version_result.data:
        raise HTTPException(status_code=404, detail=f"artifact_version {req.artifact_version_id} not found")
    version = version_result.data

    # Download raw artifact content from Supabase Storage
    storage_path = version["storage_path"]
    ext = storage_path.rsplit(".", 1)[-1] if "." in storage_path else "bin"

    try:
        raw_bytes = supabase.storage.from_("artifacts").download(storage_path)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Storage download failed: {e}")

    # Re-extract sections from the downloaded content
    payer_format = source.get("payer_format", "uhc_narrative")
    table_rows: list[list[str]] | None = None

    if ext == "pdf":
        sections, total_pages = extract_pdf_sections(raw_bytes)
        # For Priority Health: also extract raw table rows (no LLM processing)
        if payer_format == "priority_health_mdl":
            table_rows = extract_pdf_tables(raw_bytes)
        fetch_method = "pdf"
    elif ext in ("docx", "doc"):
        sections = extract_docx_sections(raw_bytes)
        total_pages = None
        fetch_method = "docx"
    elif ext in ("html", "htm"):
        sections = extract_html_sections(raw_bytes.decode("utf-8", errors="replace"), version["storage_path"])
        total_pages = None
        fetch_method = "html"
    elif ext == "csv":
        sections = extract_csv_sections(raw_bytes.decode("utf-8", errors="replace"), source.get("name", ""))
        total_pages = None
        fetch_method = "csv"
    else:
        sections = [PageSection(section_title=None, page_number=None,
                                text=raw_bytes.decode("utf-8", errors="replace")[:50000])]
        total_pages = None
        fetch_method = "api"

    # Build pipeline state — no re-fetch, content comes from Storage
    initial_state = {
        "source_id": req.source_id,
        "user_id": req.user_id,
        "payer_format": payer_format,
        "payer_name": source.get("payer_name", "Unknown Payer"),
        "artifact_version_id": req.artifact_version_id,
        "fetch_result": {
            "source_id": req.source_id,
            "content_hash": version["content_hash"],
            "storage_path": storage_path,
            "sections": [s.model_dump() for s in sections],
            "total_pages": total_pages,
            "fetch_method": fetch_method,
            "fetched_at": version.get("fetched_at", ""),
            "changed": True,  # always True here — caller already verified change
            "payer_format": payer_format,
            "payer_name": source.get("payer_name", "Unknown Payer"),
            "plan_type": source.get("plan_type", "commercial"),
            "text": "\n".join(s.text for s in sections),  # full raw text for extractors
            "table_rows": table_rows,                      # pre-parsed for Priority Health
            "source_name": source.get("name", "Unknown"),
        },
        "chunks": [],
        "extraction": None,
        "evaluation": None,
        "draft_id": None,
        "errors": [],
    }

    result = await pipeline_graph.ainvoke(initial_state)

    return {
        "status": "completed",
        "source_id": req.source_id,
        "draft_id": result.get("draft_id"),
        "eval_score": result.get("evaluation", {}).get("final_score"),
        "eval_passed": result.get("evaluation", {}).get("passed"),
        "errors": result.get("errors", []),
    }


class EmbedRequest(BaseModel):
    text: str


@app.post("/embed")
async def embed_text(req: EmbedRequest):
    """Embed a single text string using local sentence-transformers."""
    from pipeline.embedder import embed_texts
    import asyncio
    loop = asyncio.get_event_loop()
    vectors = await loop.run_in_executor(None, embed_texts, [req.text])
    return {"embedding": vectors[0]}


@app.get("/health")
async def health():
    return {"status": "ok", "service": "rxmonitor-fetcher"}
